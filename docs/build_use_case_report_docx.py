from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DEFAULT_SOURCE = ROOT / "chicktemp_use_case_report.md"
DEFAULT_OUT = ROOT / "chicktemp_use_case_report_formatted.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "000000", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_twips))


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_twips: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_twips)))

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_twips:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_twips[idx])


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def clean_text(text: str) -> str:
    text = text.replace("<br>", "\n")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def parse_report_rows(source: Path) -> tuple[list[str], list[list[str]]]:
    lines = source.read_text(encoding="utf-8").splitlines()
    table_lines = [line for line in lines if line.startswith("| ")]
    header = [part.strip() for part in table_lines[0].strip("|").split("|")]
    rows = []
    for line in table_lines[2:]:
        rows.append([clean_text(part) for part in line.strip("|").split("|")])
    return header, rows


def set_paragraph_text(cell, text: str, bold: bool = False, size: float = 7.5) -> None:
    cell.text = ""
    parts = text.split("\n")
    for idx, part in enumerate(parts):
        paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("USE CASE REPORTS")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("ChickTemp Farm Management Mobile Application")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_report_table(doc: Document, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = False

    widths = [
        1260,  # Use Case Name
        1080,  # Actors
        1800,  # Description
        1800,  # Preconditions
        2700,  # Flow of Events
        1700,  # Postconditions
        2500,  # Exceptions
    ]
    set_table_geometry(table, widths)

    header_row = table.rows[0]
    repeat_table_header(header_row)
    for idx, label in enumerate(header):
        cell = header_row.cells[idx]
        set_cell_shading(cell, "E6E6E6")
        set_cell_border(cell)
        set_cell_margins(cell, top=90, bottom=90, start=70, end=70)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_paragraph_text(cell, label, bold=True, size=7.5)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_paragraph_text(cell, value, bold=(idx == 0), size=7.0)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()

    header, rows = parse_report_rows(args.source)
    doc = Document()
    style_doc(doc)
    add_title(doc)
    add_report_table(doc, header, rows)
    doc.save(args.out)
    print(args.out)


if __name__ == "__main__":
    main()
